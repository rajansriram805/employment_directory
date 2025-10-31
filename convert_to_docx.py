#!/usr/bin/env python3
"""
Convert Markdown files to Word documents (.docx)
Requires: pip install python-docx markdown pypandoc
"""

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
    import sys
except ImportError:
    print("Installing required packages...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "markdown"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re

def markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document"""
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines (but keep some spacing)
        if not line:
            i += 1
            continue
        
        # Heading 1
        if line.startswith('# ') and not line.startswith('##'):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            i += 1
        
        # Heading 2
        elif line.startswith('## ') and not line.startswith('###'):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            i += 1
        
        # Heading 3
        elif line.startswith('### ') and not line.startswith('####'):
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
            i += 1
        
        # Heading 4
        elif line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_heading(text, level=4)
            i += 1
        
        # Code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                p.style = 'No Spacing'
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            i += 1
        
        # Horizontal rule
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)
            i += 1
        
        # Regular paragraph
        else:
            # Clean up markdown syntax
            text = line
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.+?)\*', r'\1', text)      # Italic
            text = re.sub(r'`(.+?)`', r'\1', text)       # Inline code
            text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)  # Links
            
            # Check if it's a list item
            if re.match(r'^[-*]\s+', text) or re.match(r'^\d+\.\s+', text):
                text = re.sub(r'^[-*]\s+', '', text)
                text = re.sub(r'^\d+\.\s+', '', text)
                doc.add_paragraph(text, style='List Bullet')
            else:
                doc.add_paragraph(text)
            i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✅ Successfully created: {docx_file}")

if __name__ == '__main__':
    import os
    
    files_to_convert = [
        ('PROJECT_REPORT_FOR_WORD.md', 'PROJECT_REPORT.docx'),
        ('API_DOCUMENTATION_FOR_WORD.md', 'API_DOCUMENTATION.docx')
    ]
    
    for md_file, docx_file in files_to_convert:
        if os.path.exists(md_file):
            try:
                markdown_to_docx(md_file, docx_file)
            except Exception as e:
                print(f"❌ Error converting {md_file}: {e}")
        else:
            print(f"⚠️  File not found: {md_file}")

